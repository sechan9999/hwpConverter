"""
HWP to DOCX Converter
- 1순위: CloudConvert API (CLOUDCONVERT_API_KEY 환경변수 설정 시) — 표·이미지·서식 완벽 변환
- 폴백:  pyhwp + python-docx (API 키 없을 때) — 텍스트/단락만 변환
"""

import os
import re
import time
import logging
import tempfile
import zipfile
import asyncio
import xml.etree.ElementTree as ET
from io import BytesIO
from pathlib import Path
from urllib.parse import quote

import requests as req_lib
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import Response, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document

logging.basicConfig(level=logging.INFO)
log = logging.getLogger(__name__)

BASE   = Path(__file__).parent
STATIC = BASE / "static"

MAX_MB                = int(os.environ.get("MAX_MB", 30))
MAX_ZIP_MB            = int(os.environ.get("MAX_ZIP_MB", 200))
CORS_ORIGINS          = os.environ.get("CORS_ORIGINS", "*").split(",")
CLOUDCONVERT_API_KEY  = os.environ.get("CLOUDCONVERT_API_KEY", "")
CLOUDCONVERT_API_URL  = "https://api.cloudconvert.com/v2"

app = FastAPI(title="HWP to DOCX Converter")

app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)


# ─────────────────────────────────────────
# CloudConvert 변환 (1순위)
# ─────────────────────────────────────────

def _cloudconvert(filename: str, file_bytes: bytes) -> bytes:
    """CloudConvert API v2로 HWP → DOCX 변환 (표·이미지·서식 지원)."""
    headers = {
        "Authorization": f"Bearer {CLOUDCONVERT_API_KEY}",
        "Content-Type": "application/json",
    }

    # 1. Job 생성 (upload → convert → export)
    r = req_lib.post(
        f"{CLOUDCONVERT_API_URL}/jobs",
        json={
            "tasks": {
                "upload-hwp": {"operation": "import/upload"},
                "convert-docx": {
                    "operation": "convert",
                    "input": "upload-hwp",
                    "output_format": "docx",
                },
                "export-docx": {
                    "operation": "export/url",
                    "input": "convert-docx",
                },
            }
        },
        headers=headers,
        timeout=30,
    )
    r.raise_for_status()
    job_data = r.json()["data"]
    job_id   = job_data["id"]
    log.info("CloudConvert job created: %s", job_id)

    # 2. 업로드 태스크에서 form URL/파라미터 추출
    upload_task = next(t for t in job_data["tasks"] if t["name"] == "upload-hwp")
    form        = upload_task["result"]["form"]
    upload_url  = form["url"]
    upload_params = form["parameters"]

    # 3. 파일 업로드 (multipart form)
    r = req_lib.post(
        upload_url,
        data=upload_params,
        files={"file": (filename, file_bytes, "application/octet-stream")},
        timeout=120,
    )
    r.raise_for_status()
    log.info("CloudConvert upload done")

    # 4. 완료 대기 (최대 90초 폴링)
    for attempt in range(30):
        time.sleep(3)
        r = req_lib.get(
            f"{CLOUDCONVERT_API_URL}/jobs/{job_id}",
            headers=headers,
            timeout=15,
        )
        r.raise_for_status()
        job_status = r.json()["data"]
        status     = job_status["status"]
        log.info("CloudConvert status [%d]: %s", attempt + 1, status)

        if status == "finished":
            break
        if status == "error":
            tasks_info = {t["name"]: t.get("message") for t in job_status["tasks"]}
            raise RuntimeError(f"CloudConvert 변환 오류: {tasks_info}")
    else:
        raise RuntimeError("CloudConvert 변환 시간 초과 (90초)")

    # 5. 결과 다운로드
    export_task   = next(t for t in job_status["tasks"] if t["name"] == "export-docx")
    download_url  = export_task["result"]["files"][0]["url"]
    r = req_lib.get(download_url, timeout=60)
    r.raise_for_status()
    log.info("CloudConvert download done, size=%d bytes", len(r.content))
    return r.content


# ─────────────────────────────────────────
# pyhwp 폴백 변환
# ─────────────────────────────────────────

def _iter_text_from_xml(xml_bytes: bytes):
    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return

    def local(tag):
        return tag.split("}")[-1] if "}" in tag else tag

    for elem in root.iter():
        if local(elem.tag) == "P":
            has_content = False
            for child in elem.iter():
                tag = local(child.tag)
                if tag in ("CHAR", "T") and child.text:
                    yield (child.text, False, False)
                    has_content = True
            if has_content:
                yield ("\n", False, False)


def _hwpx_fallback(hwp_bytes: bytes) -> bytes:
    doc       = Document()
    max_bytes = MAX_ZIP_MB * 1024 * 1024

    with zipfile.ZipFile(BytesIO(hwp_bytes)) as zf:
        if sum(i.file_size for i in zf.infolist()) > max_bytes:
            raise ValueError(f"압축 해제 크기가 {MAX_ZIP_MB}MB를 초과합니다.")

        section_files = sorted(
            n for n in zf.namelist()
            if re.search(r"section\d+\.xml$", n, re.IGNORECASE)
        ) or [n for n in zf.namelist() if n.endswith(".xml")]

        for sec_file in section_files:
            para = doc.add_paragraph()
            for text, bold, italic in _iter_text_from_xml(zf.read(sec_file)):
                if text == "\n":
                    para = doc.add_paragraph()
                else:
                    run       = para.add_run(text)
                    run.bold  = bold
                    run.italic = italic

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _hwp5_fallback(hwp_bytes: bytes) -> bytes:
    try:
        from hwp5.hwp5txt import Hwp5File  # type: ignore
    except ImportError as e:
        raise RuntimeError("pyhwp 패키지가 필요합니다: pip install pyhwp") from e

    doc = Document()
    with tempfile.NamedTemporaryFile(suffix=".hwp", delete=False) as tmp:
        tmp.write(hwp_bytes)
        tmp_path = tmp.name

    try:
        hwpfile = Hwp5File(tmp_path)
        for section in hwpfile.text.sections:
            for item in section.models():
                if item.get("tagname") == "HWPTAG_PARA_TEXT":
                    payload = item.get("payload", b"")
                    try:
                        text = payload.decode("utf-16-le")
                        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
                        doc.add_paragraph(text.rstrip("\r\n"))
                    except (UnicodeDecodeError, ValueError):
                        pass
        del hwpfile  # Windows: 파일 핸들 해제 후 unlink
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ─────────────────────────────────────────
# 공통 변환 진입점
# ─────────────────────────────────────────

def _convert(filename: str, file_bytes: bytes) -> tuple[bytes, str]:
    """
    Returns (docx_bytes, engine_used).
    engine_used: 'cloudconvert' | 'pyhwp-fallback'
    """
    # CloudConvert 우선
    if CLOUDCONVERT_API_KEY:
        try:
            return _cloudconvert(filename, file_bytes), "cloudconvert"
        except Exception:
            log.exception("CloudConvert 실패 — pyhwp 폴백으로 재시도")

    # pyhwp 폴백
    ext    = os.path.splitext(filename)[1].lower()
    is_zip = file_bytes[:4] == b"PK\x03\x04"
    is_ole = file_bytes[:8] == bytes.fromhex("d0cf11e0a1b11ae1")

    if ext == ".hwpx" or (ext == ".hwp" and is_zip):
        return _hwpx_fallback(file_bytes), "pyhwp-fallback"
    if ext == ".hwp" and is_ole:
        return _hwp5_fallback(file_bytes), "pyhwp-fallback"
    if ext == ".hwp":
        return _hwp5_fallback(file_bytes), "pyhwp-fallback"
    raise ValueError(f"지원하지 않는 파일 형식: {ext}")


# ─────────────────────────────────────────
# API 엔드포인트
# ─────────────────────────────────────────

@app.post("/convert")
async def convert(file: UploadFile = File(...)):
    filename = os.path.basename(file.filename or "upload.hwp").strip()
    ext      = os.path.splitext(filename)[1].lower()
    if ext not in (".hwp", ".hwpx"):
        raise HTTPException(400, "HWP 또는 HWPX 파일만 업로드 가능합니다.")

    content  = await file.read()
    size_mb  = len(content) / 1024 / 1024
    if size_mb > MAX_MB:
        raise HTTPException(413, f"파일 크기가 {MAX_MB}MB를 초과합니다. ({size_mb:.1f}MB)")

    log.info("Received %s (%.2f MB)", filename, size_mb)

    try:
        docx_bytes, engine = await asyncio.to_thread(_convert, filename, content)
    except Exception:
        log.exception("변환 실패: %s", filename)
        raise HTTPException(500, "변환 중 오류가 발생했습니다. 파일을 확인해 주세요.")

    out_name     = os.path.splitext(filename)[0] + ".docx"
    encoded_name = quote(out_name, safe="")
    log.info("Converted → %s (engine=%s)", out_name, engine)

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}",
            "X-Conversion-Engine": engine,
        },
    )


@app.get("/health")
def health():
    engine = "cloudconvert" if CLOUDCONVERT_API_KEY else "pyhwp-fallback"
    return {"status": "ok", "engine": engine}


app.mount("/", StaticFiles(directory=str(STATIC), html=True), name="static")
